from flask import Flask, request, jsonify
from langchain_community.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document as LangchainDocument
from langchain_postgres.vectorstores import PGVector
from docx import Document
from flask_cors import CORS
import google.generativeai as genai
import textwrap
import os
from config import API_KEY

app = Flask(__name__)
CORS(app, resources={r"/query": {"origins": "http://localhost:4200"}})

# Configurer l'API Generative AI de Google
genai.configure(api_key=API_KEY)

# Charger et traiter le document
def load_and_process_document(doc_path):
    doc = Document(doc_path)
    full_text = [para.text for para in doc.paragraphs]
    document_text = '\n'.join(full_text)
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=80)
    texts = text_splitter.split_text(document_text)
    return texts

# Initialiser le magasin de vecteurs
def initialize_vector_store(texts, connection_string, collection_name):
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    documents = [LangchainDocument(page_content=text) for text in texts]
    db = PGVector.from_documents(embedding=embeddings, documents=documents, 
                                 collection_name=collection_name, connection=connection_string)
    return db

# Fonction pour créer le prompt RAG
def make_rag_prompt(query, retrieved_documents):
    escaped_passage = retrieved_documents.replace("'", "").replace('"', "").replace("\n", " ")
    prompt = f"""Vous êtes un agent utile et informatif qui répond aux questions en utilisant le texte du passage de référence inclus ci-dessous.
Assurez-vous de répondre par une phrase complète, en étant exhaustif, en incluant toutes les informations de base pertinentes.
Cependant, vous parlez à un public non technique, alors assurez-vous de décomposer les concepts compliqués et
adoptez un ton amical et conversationnel.
QUESTION : '{query}'
PASSAGE : '{escaped_passage}'

RÉPONSE :
"""
    return prompt

# Fonction pour générer une réponse à partir du modèle Generative AI
def generate_answer(prompt):
    model = genai.GenerativeModel("gemini-1.5-flash")
    result = model.generate_content(prompt)
    return result.text

# Charger et traiter le document
texts = load_and_process_document('Guidecrédit.docx')

# Initialiser le magasin de vecteurs
CONNECTION_STRING = "postgresql+psycopg2://postgres:123@localhost:5432/vectorDB"
COLLECTION_NAME = "status_of_doc_vectors"
db = initialize_vector_store(texts, CONNECTION_STRING, COLLECTION_NAME)

@app.route('/query', methods=['POST'])
def query_service():
    data = request.json
    query = data.get('query')
    
    if not query:
        return jsonify({"error": "Query not provided"}), 400

    try:
        results = db.similarity_search_with_score(query, k=5)
        retrieved_documents = "".join([doc[0].page_content for doc in results])
        final_prompt = make_rag_prompt(query, retrieved_documents)
        answer = generate_answer(final_prompt)
        return jsonify({"answer": answer})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8888)
